"""
generate_cv.py — ATS-optimised Word resume for Wesley Gomes da Silva.
Output: ../src/app/assets/WesleySilva_Latam_AgileProjectDelivery_BR_C1English.docx
        ../public/WesleySilva_Latam_AgileProjectDelivery_BR_C1English.docx
Run  : python generate_cv.py
Deps : pip install python-docx
"""

from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


CV_FILENAME = "WesleySilva_Latam_AgileProjectDelivery_BR_C1English.docx"
OUTPUT      = Path(__file__).resolve().parent.parent / "src" / "app" / "assets" / CV_FILENAME
OUTPUT_PUB  = Path(__file__).resolve().parent.parent / "public" / CV_FILENAME


# ── helpers ──────────────────────────────────────────────────────────────────

def _add_hyperlink(paragraph, url, text, size=10, color=(0, 82, 155)):
    """Insert a clickable hyperlink into an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(*color))
    rPr.append(color_el)
    sz_el = OxmlElement('w:sz')
    sz_el.set(qn('w:val'), str(size * 2))
    rPr.append(sz_el)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _para(doc, text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def _section_title(doc, title):
    """Bold uppercase section header with bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def _job_header(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(title)
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(f"  |  {company}  |  {location}  |  {period}")
    r2.bold = False
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(80, 80, 80)


def _tech_stack(doc, stack):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run("Tech Stack: ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(stack)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(60, 60, 60)


# ── document ─────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins (narrow for more content space)
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.85)
        section.right_margin  = Inches(0.85)

    # Default paragraph font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # ── NAME & CONTACT ──────────────────────────────────────────────────────
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    r = name_p.add_run("Wesley Gomes da Silva")
    r.bold = True
    r.font.size = Pt(18)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    rt = title_p.add_run(
        "Agile Delivery Manager  |  AI-Driven Software Delivery  |  Tech Heavy  |  Digital Products & Distributed Teams"
    )
    rt.font.size = Pt(11)
    rt.font.color.rgb = RGBColor(60, 60, 60)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(2)
    _add_hyperlink(contact_p, 'https://linkedin.com/in/wesleyzilva', 'linkedin.com/in/wesleyzilva')
    sep1 = contact_p.add_run('  |  wesley.zilva@gmail.com  |  São Carlos, SP – Brazil')
    sep1.font.size = Pt(10)

    portfolio_p = doc.add_paragraph()
    portfolio_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    portfolio_p.paragraph_format.space_after = Pt(6)
    rp1 = portfolio_p.add_run('Portfolio: ')
    rp1.font.size = Pt(10)
    _add_hyperlink(portfolio_p, 'https://wesleyzilva.github.io/portfolioNearshoreWesIA/', 'wesleyzilva.github.io/portfolioNearshoreWesIA')
    sep3 = portfolio_p.add_run('  |  GitHub: ')
    sep3.font.size = Pt(10)
    _add_hyperlink(portfolio_p, 'https://github.com/wesleyzilva', 'github.com/wesleyzilva')

    # ── PROFESSIONAL SUMMARY ────────────────────────────────────────────────
    _section_title(doc, "Professional Summary")
    _para(doc,
        "10+ years of experience building mission-critical systems and leading distributed, "
        "high-performance teams at global companies such as Experian and Amdocs. "
        "Specialised in delivery of remote and nearshore engineering squads across Brazil, India, "
        "and South Africa, connecting business strategy with technical execution to deliver scalable "
        "digital products powered by Observability, DevSecOps, and AI-driven delivery. "
        "Leveraging GitHub Copilot and AI-driven workspaces to streamline workflows and accelerate "
        "the SDLC by up to 45%. Led data pipelines (Databricks, Spark) supporting reconciliation of "
        "R$500M+ monthly. Reduced critical vulnerabilities by 93% ensuring compliance with PCI-DSS and LGPD.",
        size=10, space_after=4)

    # ── CORE COMPETENCIES ───────────────────────────────────────────────────
    _section_title(doc, "Core Competencies")
    competencies = [
        "Global & Nearshore Team Leadership  •  Agile Project Management (SAFe / Scrum / Kanban)  •  OKRs & Roadmaps",
        "Data Engineering at Scale (Databricks / Spark / ETL)  •  DevSecOps & Compliance (PCI-DSS / LGPD / BACEN / FEBRABAN)",
        "System Observability (Datadog / Dynatrace)  •  CI/CD Pipelines  •  DORA Metrics  •  AI-Powered Workflows",
        "Full-Stack Development (Java / Angular / .NET / OutSystems)  •  Security (Okta / AWS Security / Veracode / Rapid7)",
    ]
    for c in competencies:
        _para(doc, c, size=10, space_after=2)

    # ── TECHNICAL SKILLS ────────────────────────────────────────────────────
    _section_title(doc, "Technical Skills")
    skills_rows = [
        ("Data Engineering",   "Databricks, Apache Spark, Pentaho ETL, Python, SQL Advanced, APIs REST"),
        ("Development",        "Java Spring, .NET, Angular, TypeScript, JavaScript, OutSystems, Salesforce"),
        ("Security",           "Okta, AWS Security, Veracode, Rapid7, PCI-DSS, LGPD, BACEN, FEBRABAN"),
        ("DevOps & CI/CD",     "GitHub Actions, Jenkins, Docker, Kubernetes, Linux, Shell Scripting"),
        ("Observability",      "Datadog (Certified — Dashboards, Metrics, Tagging), Dynatrace, APM, RUM"),
        ("Databases",          "PostgreSQL, SQL Server, OracleDB, PL/SQL"),
        ("Agile Tools",        "JIRA, Azure Boards, Confluence, Figma, ServiceNow"),
        ("AI Tools",           "GitHub Copilot, ChatGPT, Datadog AI"),
        ("Languages",          "English (Native/Bilingual), Portuguese (Native), Spanish (Elementary), Chinese (Elementary)"),
    ]
    for label, items in skills_rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(items)
        r2.font.size = Pt(10)

    # ── 10 YEARS AT SERASA EXPERIAN ─────────────────────────────────────────
    _section_title(doc, "Core Experience — 10 Years at Serasa Experian")

    _para(doc,
        "Serasa Experian is Brazil's leading credit bureau and data analytics company (Experian Group), "
        "processing 80M+ daily financial transactions across fintech and payment systems.",
        size=10, space_after=2)

    # ---- Role 1: IT Coordinator ----
    _job_header(doc,
        "IT Coordinator — M&A Startup Fintech Flexpag | Payments & Data",
        "Serasa Experian",
        "Remote · São Carlos / Blumenau / Recife",
        "Nov 2024 – Present (1 yr 7 mos)")
    _para(doc, "Led Data & ETL squads (2 squads × 4 people) in M&A fintech integration.", size=10, space_after=2)
    for b in [
        "Architecture Optimisation: Led Data squads in modernising architecture with Databricks, ensuring higher processing robustness and real-time, data-driven decision-making.",
        "Delivery Acceleration: Significantly reduced time-to-market for analytical products by streamlining CI/CD pipelines and release workflows.",
        "Payments & ETL: Managed the ETL squad focused on complex acquiring and sub-acquiring financial processes via Pentaho and REST APIs, ensuring 100% precision in billing and revenue reconciliation.",
        "Regulatory Compliance: Guaranteed strict adherence to BACEN and FEBRABAN regulatory reporting requirements.",
        "DevSecOps Zero-to-One: Fully implemented the DevSecOps framework and vulnerability management programme using Veracode and Rapid7 from the ground up — reducing critical vulnerabilities by 93%.",
        "SDLC Security: Integrated security layers throughout the SDLC, hardening infrastructure and sensitive data under PCI-DSS standards.",
    ]:
        _bullet(doc, b)
    _tech_stack(doc, "Databricks, Apache Spark, Pentaho ETL, Python, Veracode, Rapid7, Jenkins, GitHub Actions, Datadog, PCI-DSS, LGPD, JIRA")

    # ---- Role 2: Agile Delivery Manager ----
    _job_header(doc,
        "Agile Delivery Manager — HR, Legal, Customer Service, Sales & Billing",
        "Serasa Experian",
        "São Carlos, Brazil (Hybrid)",
        "Apr 2021 – Nov 2024 (3 yrs 8 mos)")
    _para(doc, "Led 5 squads × 7 people. Stakeholders aligned with London and USA.", size=10, space_after=2)
    for b in [
        "Delivery Metrics Management: Implemented and monitored KPIs and efficiency metrics (Velocity, Burndown, Cycle Time), generating data-driven insights for continuous team performance improvement.",
        "SDLC Optimisation: Drove Agile methodology adoption across Technology and Data squads, reducing development bottlenecks and increasing productivity through optimised workflows.",
        "Governance & Roadmaps: Built strategic roadmaps for scalable solutions, ensuring alignment between OKRs and engineering delivery capacity across 4 business units.",
        "Design-Engineering Interface: Proactive Figma collaboration to ensure technical fidelity during prototype-to-front-end handoff, reducing rework by 40%.",
        "Cross-BU Synergy: Managed technical interdependencies across Legal, Finance, Sales — ensuring systemic integrity and continuous value delivery in hybrid environments.",
        "AI-Powered Engineering: Leveraged GitHub Copilot and AI-driven workspaces to streamline workflows and accelerate the SDLC by up to 45%.",
    ]:
        _bullet(doc, b)
    _tech_stack(doc, "SAFe, Scrum, Kanban, JIRA, Azure Boards, Confluence, Figma, DORA Metrics, ServiceNow, GitHub Copilot")

    # ---- Role 3: Full Stack Developer ----
    _job_header(doc,
        "Full Stack Developer — Billing, ITIL & Observability",
        "Serasa Experian",
        "São Carlos, Brazil",
        "Mar 2016 – Apr 2021 (5 yrs 2 mos)")
    for b in [
        "End-to-End Development: Acted across the full development lifecycle of core Billing and CRM systems using Java and Angular, delivering scalable, high-availability solutions.",
        "Observability Leadership: Technical lead for mission-critical system support using Dynatrace and Datadog (APM + RUM), reducing incident MTTR by 45% and repeat incidents by 60%.",
        "Architecture & Requirements: Translated business needs into high-fidelity technical specifications aligned with corporate SLAs.",
        "Technical Ecosystem: Integrated and evolved Salesforce, OutSystems, and .NET platforms, navigating between legacy and modern technologies.",
        "Data Engineering & Backend: Designed and maintained complex databases, ensuring data integrity for high volumes of financial transactions.",
    ]:
        _bullet(doc, b)
    _tech_stack(doc, "Java Spring, .NET, Angular, TypeScript, Dynatrace, Datadog, APM, RUM, OutSystems, Salesforce, SQL Server, PostgreSQL")

    # ── OTHER EXPERIENCE ─────────────────────────────────────────────────────
    _section_title(doc, "Other Experience")

    # ---- Amdocs ----
    _job_header(doc,
        "Technical Support Specialist",
        "Amdocs Studios",
        "São Carlos, Brazil",
        "Jul 2013 – Mar 2016 (2 yrs 9 mos)")
    for b in [
        "Operational Stability & Automation: Deep expertise in Unix environments, developing Shell Script and PL/SQL routines to optimise OracleDB performance — stable environments and proactive failure reduction.",
        "Incident Management with India/USA Shifts: Proficiency in Remedy and CRM Clarify 12 for agile relationship platform support, directly impacting end-user satisfaction.",
        "Critical Middleware Support: Administration of Tuxedo and Weblogic, ensuring seamless transactional performance between essential systems.",
    ]:
        _bullet(doc, b)
    _tech_stack(doc, "Unix (AIX/Solaris), OracleDB, PL/SQL, Shell Scripting, Tuxedo, Weblogic, Remedy, CRM Clarify 12")

    # ---- XBot ----
    _job_header(doc,
        "Business Development Manager & Programming Language Professor",
        "XBot",
        "São Carlos, Brazil",
        "Nov 2010 – Jul 2013 (2 yrs 9 mos)")
    for b in [
        "Managed 12 sales and partner teams across Brazil focused on technological education with robots, energy systems, and security/simulation systems.",
        "Served as Programming Language Professor (Jan–Dec 2011).",
    ]:
        _bullet(doc, b)

    # ---- Earlier roles (condensed) ----
    _job_header(doc,
        "Earlier Roles",
        "Bernardin Consultants (FR) · Sidertec · Elektro · Tecumseh Products",
        "Brazil / France",
        "2005 – 2010")
    for b in [
        "International Consultant at Bernardin Consultants (France Contractor, 2010): Market development for high-tech French/Swiss equipment (infrared microchips, precision dosing machines, aerospace connectors) — bridge between foreign investors and the Brazilian market.",
        "IT Analyst at Sidertec Estruturas Metálicas (2009–2010): End-to-end IT infrastructure administration.",
        "Trainee Maintenance Engineering at Elektro (2009): Maintenance planning (SAP PM), Scada Eclipse, NR10 safety instructor, IEC 61850 substation automation.",
        "Trainee Engineering at Tecumseh Products (2005–2008): Materials specification, supplier qualification, preventive/predictive maintenance scheduling, Oracle ERP administration.",
    ]:
        _bullet(doc, b)

    # ── CERTIFICATIONS ──────────────────────────────────────────────────────
    _section_title(doc, "Certifications")
    certs = [
        ("Certified Angular Developer",        "Angular Certification Body"),
        ("Certified JavaScript Developer",      "JavaScript Certification Body"),
        ("Datadog — Dashboards",                "Datadog"),
        ("Datadog — Metrics",                   "Datadog"),
        ("Datadog — Tagging Best Practices",    "Datadog"),
        ("Agile Coach Recognition",             "Serasa Experian CX & Tech"),
        ("Yellow Belt Six Sigma",               "Process Improvement & Lean"),
    ]
    for name, issuer in certs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(f"{name}  —  ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(issuer)
        r2.font.size = Pt(10)

    # ── EDUCATION ───────────────────────────────────────────────────────────
    _section_title(doc, "Education")
    edu = [
        ("Bachelor's in Engineering (Recloser Retrofit & Power Transmission Systems)",
         "Centro Universitário Central Paulista", "2005 – 2011"),
        ("Graduate Certificate in Web Development — Java and Databases Specialisation",
         "Universidade Federal de São Carlos (UFSCar)", "2018"),
        ("Electrical Maintenance Technician",
         "SENAI São Paulo", "1999 – 2000"),
    ]
    for degree, institution, period in edu:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{degree}  —  ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(f"{institution}  |  {period}")
        r2.font.size = Pt(10)

    # ── PUBLICATIONS ────────────────────────────────────────────────────────
    _section_title(doc, "Publications")
    pubs = [
        "A block programming interface for educational mobile robots",
        "Retrofit of reclosers in electric power grids",
    ]
    for pub in pubs:
        _bullet(doc, pub)

    # ── KEY METRICS ─────────────────────────────────────────────────────────
    _section_title(doc, "Key Metrics & Achievements")
    metrics = [
        "93% reduction in critical vulnerabilities — zero BACEN/FEBRABAN audit findings",
        "R$500M+ monthly reconciliation with 99.8% accuracy across 80M+ daily transactions",
        "45% SDLC acceleration via AI-driven engineering workflows (GitHub Copilot)",
        "40% ETL processing time reduction via Databricks architecture optimisation",
        "25% productivity increase across nearshore squads in 3 time zones (Brazil / India / South Africa)",
        "35% sprint velocity increase across 5+ squads — 4 business units",
        "40% design-to-code rework reduction via Figma-to-dev workflow alignment",
        "45% MTTR reduction — 99.95% uptime in mission-critical systems",
    ]
    for m in metrics:
        _bullet(doc, m)

    # ── SAVE ─────────────────────────────────────────────────────────────────
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"✓  CV saved → {OUTPUT}")

    OUTPUT_PUB.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(OUTPUT, OUTPUT_PUB)
    print(f"✓  CV copied → {OUTPUT_PUB}")


if __name__ == "__main__":
    build()
